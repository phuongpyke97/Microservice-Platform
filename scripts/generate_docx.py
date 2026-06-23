import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def add_heading_styled(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Segoe UI'
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    return h

def main():
    doc = docx.Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark Gray

    # Colors
    NAVY = RGBColor(0x1B, 0x36, 0x5D)
    SLATE = RGBColor(0x5C, 0x76, 0x8D)
    RED = RGBColor(0xC0, 0x39, 0x2B)

    # Title Page / Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("HƯỚNG DẪN TÍCH HỢP API CMS ADMIN")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Epic 3: Quản lý Kho Nhạc Hệ Thống (Music Library Management)")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = SLATE
    subtitle.paragraph_format.space_after = Pt(24)

    # Introduction
    p = doc.add_paragraph()
    p.add_run("Tài liệu này cung cấp chi tiết các API đầu cuối và hướng dẫn tích hợp giao diện cho lập trình viên ").font.name = 'Segoe UI'
    p.add_run("Frontend (FE) ").bold = True
    p.add_run("khi phát triển trang Quản lý Kho Nhạc trên CMS Admin. Hệ thống sử dụng mô hình tải tệp trực tiếp lên bộ lưu trữ đám mây qua ").font.name = 'Segoe UI'
    p.add_run("Presigned URL ").bold = True
    p.add_run("nhằm tối ưu hóa băng thông mạng và hiệu năng máy chủ.")

    # 1. Host và Thông tin chung
    add_heading_styled(doc, "1. Thông tin kết nối chung", 1, NAVY)
    
    p_host = doc.add_paragraph()
    p_host.add_run("• API Gateway Endpoint: ").bold = True
    p_host.add_run("http://103.154.62.118:18080 (hoặc http://localhost:18080)\n")
    p_host.add_run("• Base Path dịch vụ File (Upload): ").bold = True
    p_host.add_run("/file-service/api/files\n")
    p_host.add_run("• Base Path dịch vụ Nhạc (Library): ").bold = True
    p_host.add_run("/crbt-community-library/library\n")
    p_host.add_run("• Các Bucket MinIO sử dụng: ").bold = True
    p_host.add_run("media-temp (tạm thời), media-audio (nhạc cá nhân), media-audio-lib (thư viện nhạc nền admin), media-images (ảnh bìa)\n")
    p_host.add_run("• Định dạng dữ liệu mặc định: ").bold = True
    p_host.add_run("JSON (application/json)")

    # 2. Luồng tải file (Upload)
    add_heading_styled(doc, "2. Luồng nghiệp vụ Tải nhạc & Xác thực dữ liệu", 1, NAVY)
    
    p_flow = doc.add_paragraph()
    p_flow.add_run("Để upload một bài nhạc mới lên hệ thống, FE cần thực hiện theo luồng 4 bước sau:\n\n")
    
    p_flow.add_run("Bước 1: Lấy Presigned URL tải lên\n").bold = True
    p_flow.add_run("FE gửi request lên API gateway để đăng ký file và nhận Presigned URL:\n")
    p_flow.add_run("• API: ").italic = True
    p_flow.add_run("GET /file-service/api/files/presigned/upload?originalName=ten_file.mp3&contentType=audio/mpeg\n")
    p_flow.add_run("• Response nhận về gồm: ").italic = True
    p_flow.add_run("fileId (ID file trong DB), objectKey, và url (URL ký sẵn của MinIO - đã được map với IP ngoại/localhost thực tế).\n\n")

    p_flow.add_run("Bước 2: Upload trực tiếp file lên MinIO\n").bold = True
    p_flow.add_run("FE thực hiện upload trực tiếp file âm thanh (MP3/WAV) từ trình duyệt lên MinIO bằng URL ký sẵn nhận ở Bước 1:\n")
    p_flow.add_run("• Method: ").italic = True
    p_flow.add_run("PUT\n")
    p_flow.add_run("• URL: ").italic = True
    p_flow.add_run("Sử dụng chính xác giá trị 'url' nhận ở Bước 1.\n")
    p_flow.add_run("• Header: ").italic = True
    p_flow.add_run("Content-Type đặt là 'audio/mpeg' (MP3) hoặc 'audio/wav' (WAV) tương ứng. (Lưu ý chọn Authorization = No Auth).\n")
    p_flow.add_run("• Body: ").italic = True
    p_flow.add_run("Chứa file binary nhị phân của bài nhạc.\n\n")

    p_flow.add_run("Bước 3: Xác nhận tải file thành công (Confirm file)\n").bold = True
    p_flow.add_run("Sau khi MinIO trả về 200 OK, FE gọi API confirm để hệ thống chuyển file từ thư mục tạm sang thư mục lưu trữ chính thức:\n")
    p_flow.add_run("• API: ").italic = True
    p_flow.add_run("POST /file-service/api/files/{fileId}/confirm\n")
    p_flow.add_run("• Body: ").italic = True
    p_flow.add_run("{\"targetBucket\": \"media-audio-lib\"} (hoặc \"media-audio\" đối với nhạc cá nhân của user)\n")
    p_flow.add_run("• Response trả về chứa: ").italic = True
    p_flow.add_run("đường dẫn file cuối cùng ('storedKey').\n\n")

    p_flow.add_run("Bước 4: Lưu thông tin bài nhạc vào kho nhạc (Save Ringtone)\n").bold = True
    p_flow.add_run("FE gọi API tạo mới bài nhạc của thư viện nhạc hệ thống kèm các trường thông tin metadata:\n")
    p_flow.add_run("• API: ").italic = True
    p_flow.add_run("POST /crbt-community-library/library/ringtones\n")
    p_flow.add_run("• Body: ").italic = True
    p_flow.add_run("{\"title\": \"Tên bài nhạc\", \"artistName\": \"Ca sĩ/Tác giả\", \"audioUrl\": \"Đường dẫn nhạc từ bước 3 hoặc full URL MinIO ngoại\", \"moodId\": 3, \"status\": true, \"categoryId\": 1}\n")

    # 3. Các ràng buộc dữ liệu (Validation)
    add_heading_styled(doc, "3. Ràng buộc dữ liệu & Mã lỗi cần bắt (Validation & Error Codes)", 1, NAVY)
    p_val = doc.add_paragraph()
    p_val.add_run("Khi FE thực hiện lưu bài nhạc ở Bước 4, hệ thống backend sẽ tự động tải file để kiểm tra các tiêu chí sau. Nếu vi phạm, hệ thống trả về mã lỗi ").font.name = 'Segoe UI'
    p_val.add_run("400 Bad Request").bold = True
    p_val.add_run(" kèm thông điệp lỗi chi tiết:\n\n")

    p_val.add_run("1. Kiểm tra Giọng hát/Lời nhạc (Vocal detection):\n").bold = True
    p_val.add_run("• Quy tắc: ").italic = True
    p_val.add_run("Chỉ chấp nhận nhạc không lời. Nếu phát hiện có giọng hát, hệ thống báo lỗi:\n")
    p_val.add_run("• Thông điệp lỗi: ").italic = True
    p_val.add_run("\"Nhạc có lời không được phép sử dụng. Vui lòng tải lên nhạc không lời.\"\n\n").font.color.rgb = RED

    p_val.add_run("2. Giới hạn dung lượng file (< 50MB):\n").bold = True
    p_val.add_run("• Quy tắc: ").italic = True
    p_val.add_run("Dung lượng file tải lên không được vượt quá 50MB.\n")
    p_val.add_run("• Thông điệp lỗi: ").italic = True
    p_val.add_run("\"Dung lượng file vượt quá giới hạn cho phép (50MB).\"\n\n").font.color.rgb = RED

    p_val.add_run("3. Giới hạn thời lượng (< 5 phút):\n").bold = True
    p_val.add_run("• Quy tắc: ").italic = True
    p_val.add_run("Độ dài bài nhạc phải ngắn hơn 5 phút (300 giây).\n")
    p_val.add_run("• Thông điệp lỗi: ").italic = True
    p_val.add_run("\"Thời lượng bài nhạc không được vượt quá 5 phút.\"").font.color.rgb = RED

    # 4. Danh sách các API cụ thể cho CMS Admin
    add_heading_styled(doc, "4. Chi tiết các API quản trị kho nhạc", 1, NAVY)

    # API 1: Statistics
    add_heading_styled(doc, "API 1: Xem Thống kê Tổng quan (Thống kê đầu trang)", 2, SLATE)
    p_api1 = doc.add_paragraph()
    p_api1.add_run("• Method / Path: ").bold = True
    p_api1.add_run("GET /library/ringtones/statistics\n")
    p_api1.add_run("• Mô tả: ").bold = True
    p_api1.add_run("Lấy thông tin hiển thị lên 4 thẻ thống kê ở đầu trang CMS.\n")
    p_api1.add_run("• Response Sample (200 OK):\n").bold = True
    p_api1_code = doc.add_paragraph()
    p_api1_code.paragraph_format.left_indent = Inches(0.5)
    r1 = p_api1_code.add_run(
        "{\n"
        "  \"code\": 0,\n"
        "  \"message\": \"Success\",\n"
        "  \"data\": {\n"
        "    \"totalTracks\": 150,        // Tổng số bài nhạc\n"
        "    \"activeTracks\": 120,       // Số bài đang hiển thị (switch bật)\n"
        "    \"inactiveTracks\": 30,       // Số bài đang ẩn (switch tắt)\n"
        "    \"totalSelections\": 4500     // Tổng lượt được chọn của toàn hệ thống\n"
        "  }\n"
        "}"
    )
    r1.font.name = 'Consolas'
    r1.font.size = Pt(9.5)

    # API 2: Search & Filter
    add_heading_styled(doc, "API 2: Tìm kiếm & Lọc Danh sách (Phân trang)", 2, SLATE)
    p_api2 = doc.add_paragraph()
    p_api2.add_run("• Method / Path: ").bold = True
    p_api2.add_run("GET /crbt-community-library/library/ringtones/search\n")
    p_api2.add_run("• Query Parameters:\n").bold = True
    p_api2.add_run("  - q (string): Tìm kiếm real-time theo Tên bài nhạc hoặc Tên ca sĩ (so khớp tương đối).\n")
    p_api2.add_run("  - categoryId (long): Lọc theo Thể loại nhạc.\n")
    p_api2.add_run("  - moodId (long): Lọc theo Tâm trạng nhạc.\n")
    p_api2.add_run("  - status (boolean): Lọc theo trạng thái (true = Hiển thị, false = Ẩn).\n")
    p_api2.add_run("  - createdFrom (string, format dd/MM/yyyy hoặc yyyy-MM-dd): Khoảng ngày tải lên (Từ ngày).\n")
    p_api2.add_run("  - createdTo (string, format dd/MM/yyyy hoặc yyyy-MM-dd): Khoảng ngày tải lên (Đến ngày).\n")
    p_api2.add_run("  - selectionCountFrom (int): Lọc dải lượt chọn (Từ số lượt).\n")
    p_api2.add_run("  - selectionCountTo (int): Lọc dải lượt chọn (Đến số lượt).\n")
    p_api2.add_run("  - page (int, mặc định 0): Chỉ số trang kết quả.\n")
    p_api2.add_run("  - size (int, mặc định 20): Số bản ghi trên một trang.\n")
    p_api2.add_run("• Response Sample (200 OK):\n").bold = True
    
    p_api2_code = doc.add_paragraph()
    p_api2_code.paragraph_format.left_indent = Inches(0.5)
    r2 = p_api2_code.add_run(
        "{\n"
        "  \"success\": true,\n"
        "  \"message\": \"OK\",\n"
        "  \"data\": {\n"
        "    \"content\": [\n"
        "      {\n"
        "        \"id\": 6,\n"
        "        \"title\": \"Warm Morning Acoustic\",\n"
        "        \"artistName\": \"Sunny Vibes\",\n"
        "        \"audioUrl\": \"http://103.154.62.118:9000/media-audio-lib/a33cb5e5-109a-4c41-ba39-9c4fbfa09666-zoluushka-guitar-type-beat.mp3\",\n"
        "        \"coverImageUrl\": \"http://localhost:9000/media-images/cover.jpg\",\n"
        "        \"durationSeconds\": 173,\n"
        "        \"featured\": true,\n"
        "        \"status\": true,\n"
        "        \"selectionCount\": 0,\n"
        "        \"category\": {\n"
        "          \"id\": 1,\n"
        "          \"name\": \"Pop\",\n"
        "          \"description\": \"Popular mainstream music\",\n"
        "          \"createdAt\": \"2026-05-25T13:23:33.108Z\",\n"
        "          \"updatedAt\": null\n"
        "        },\n"
        "        \"mood\": {\n"
        "          \"id\": 3,\n"
        "          \"name\": \"Chill\",\n"
        "          \"description\": \"Thư giãn, nhẹ nhàng, êm dịu\",\n"
        "          \"createdAt\": \"2026-05-27T15:15:21Z\",\n"
        "          \"updatedAt\": null\n"
        "        },\n"
        "        \"createdAt\": \"2026-05-27T16:11:19Z\",\n"
        "        \"updatedAt\": \"2026-05-27T16:11:19Z\"\n"
        "      }\n"
        "    ],\n"
        "    \"pageNumber\": 0,\n"
        "    \"pageSize\": 20,\n"
        "    \"totalElements\": 1,\n"
        "    \"totalPages\": 1,\n"
        "    \"last\": true\n"
        "  }\n"
        "}"
    )
    r2.font.name = 'Consolas'
    r2.font.size = Pt(9.5)

    # API 3: Export
    add_heading_styled(doc, "API 3: Xuất báo cáo CSV (Export)", 2, SLATE)
    p_api3 = doc.add_paragraph()
    p_api3.add_run("• Method / Path: ").bold = True
    p_api3.add_run("GET /crbt-community-library/library/ringtones/export\n")
    p_api3.add_run("• Query Parameters: ").bold = True
    p_api3.add_run("Giống hệt API tìm kiếm (để xuất đúng tập dữ liệu đang lọc) nhưng không truyền tham số phân trang 'page' và 'size'.\n")
    p_api3.add_run("• Tác dụng: ").bold = True
    p_api3.add_run("Tự động tải xuống file .csv chứa toàn bộ danh sách đang lọc, hiển thị tiếng Việt chuẩn trên Excel nhờ mã UTF-8 BOM.")

    # API 4: Create
    add_heading_styled(doc, "API 4: Thêm nhạc mới vào kho", 2, SLATE)
    p_api4 = doc.add_paragraph()
    p_api4.add_run("• Method / Path: ").bold = True
    p_api4.add_run("POST /crbt-community-library/library/ringtones\n")
    p_api4.add_run("• Request Body:\n").bold = True
    
    p_api4_code = doc.add_paragraph()
    p_api4_code.paragraph_format.left_indent = Inches(0.5)
    r4 = p_api4_code.add_run(
        "{\n"
        "  \"title\": \"Warm Morning Acoustic\",\n"
        "  \"artistName\": \"Sunny Vibes\",\n"
        "  \"audioUrl\": \"http://103.154.62.118:9000/media-audio-lib/a33cb5e5-109a-4c41-ba39-9c4fbfa09666-guitar.mp3\", // storedKey\n"
        "  \"coverImageUrl\": \"http://localhost:9000/media-images/cover.jpg\",\n"
        "  \"moodId\": 3, // ID Tâm trạng lấy từ API danh mục Mood (Ví dụ: 3 = Chill)\n"
        "  \"status\": true,\n"
        "  \"categoryId\": 1 // ID Thể loại lấy từ API danh mục Category (Ví dụ: 1 = Pop)\n"
        "}"
    )
    r4.font.name = 'Consolas'
    r4.font.size = Pt(9.5)

    # API 5: Edit
    add_heading_styled(doc, "API 5: Chỉnh sửa thông tin bài nhạc (Metadata)", 2, SLATE)
    p_api5 = doc.add_paragraph()
    p_api5.add_run("• Method / Path: ").bold = True
    p_api5.add_run("PUT /crbt-community-library/library/ringtones/{id}\n")
    p_api5.add_run("• Request Body:\n").bold = True
    
    p_api5_code = doc.add_paragraph()
    p_api5_code.paragraph_format.left_indent = Inches(0.5)
    r5 = p_api5_code.add_run(
        "{\n"
        "  \"title\": \"Warm Morning Acoustic (Guitar Edit)\",\n"
        "  \"artistName\": \"Sunny Vibes\",\n"
        "  \"moodId\": 3,\n"
        "  \"status\": true,\n"
        "  \"categoryId\": 1\n"
        "}"
    )
    r5.font.name = 'Consolas'
    r5.font.size = Pt(9.5)

    # API 6: Toggle Status
    add_heading_styled(doc, "API 6: Bật/Tắt Trạng thái nhanh (Inline Toggle)", 2, SLATE)
    p_api6 = doc.add_paragraph()
    p_api6.add_run("• Method / Path: ").bold = True
    p_api6.add_run("PATCH /crbt-community-library/library/ringtones/{id}/status\n")
    p_api6.add_run("• Request Body:\n").bold = True
    
    p_api6_code = doc.add_paragraph()
    p_api6_code.paragraph_format.left_indent = Inches(0.5)
    r6 = p_api6_code.add_run(
        "{\n"
        "  \"status\": false // bật = true (Hiển thị), tắt = false (Ẩn)\n"
        "}"
    )
    r6.font.name = 'Consolas'
    r6.font.size = Pt(9.5)

    # API 7: Delete
    add_heading_styled(doc, "API 7: Xóa bài nhạc khỏi kho quản lý", 2, SLATE)
    p_api7 = doc.add_paragraph()
    p_api7.add_run("• Method / Path: ").bold = True
    p_api7.add_run("DELETE /crbt-community-library/library/ringtones/{id}\n")
    p_api7.add_run("• Tác dụng: ").bold = True
    p_api7.add_run("Xóa hoàn toàn thông tin bài nhạc khỏi kho quản lý và gửi tín hiệu xóa file vật lý trên MinIO. Lượt sử dụng thống kê được sao lưu an toàn tự động ở phía DB để chạy báo cáo sau này.\n\n")

    # API 8: Categories Management
    add_heading_styled(doc, "API 8: Quản lý Thể loại (Categories CRUD)", 2, SLATE)
    p_api8 = doc.add_paragraph()
    p_api8.add_run("1. Lấy danh sách thể loại:\n").bold = True
    p_api8.add_run("• Method / Path: ").italic = True
    p_api8.add_run("GET /crbt-community-library/library/categories\n")
    p_api8.add_run("2. Thêm mới thể loại:\n").bold = True
    p_api8.add_run("• Method / Path: ").italic = True
    p_api8.add_run("POST /crbt-community-library/library/categories\n")
    p_api8.add_run("• Body: ").italic = True
    p_api8.add_run("{\"name\": \"Thể loại mới\", \"description\": \"Mô tả\"}\n")
    p_api8.add_run("3. Cập nhật thể loại:\n").bold = True
    p_api8.add_run("• Method / Path: ").italic = True
    p_api8.add_run("PUT /crbt-community-library/library/categories/{id}\n")
    p_api8.add_run("• Body: ").italic = True
    p_api8.add_run("{\"name\": \"Tên cập nhật\", \"description\": \"Mô tả mới\"}\n")
    p_api8.add_run("4. Xóa thể loại:\n").bold = True
    p_api8.add_run("• Method / Path: ").italic = True
    p_api8.add_run("DELETE /crbt-community-library/library/categories/{id}\n")
    p_api8.add_run("• Lưu ý: ").bold = True
    p_api8.add_run("Không cho phép xóa nếu danh mục này đang có bài nhạc sử dụng (trả về lỗi validation).\n\n")

    # API 9: Moods Management
    add_heading_styled(doc, "API 9: Quản lý Tâm trạng (Moods CRUD)", 2, SLATE)
    p_api9 = doc.add_paragraph()
    p_api9.add_run("1. Lấy danh sách tâm trạng:\n").bold = True
    p_api9.add_run("• Method / Path: ").italic = True
    p_api9.add_run("GET /crbt-community-library/library/moods\n")
    p_api9.add_run("2. Thêm mới tâm trạng:\n").bold = True
    p_api9.add_run("• Method / Path: ").italic = True
    p_api9.add_run("POST /crbt-community-library/library/moods\n")
    p_api9.add_run("• Body: ").italic = True
    p_api9.add_run("{\"name\": \"Tâm trạng mới\", \"description\": \"Mô tả\"}\n")
    p_api9.add_run("3. Cập nhật tâm trạng:\n").bold = True
    p_api9.add_run("• Method / Path: ").italic = True
    p_api9.add_run("PUT /crbt-community-library/library/moods/{id}\n")
    p_api9.add_run("• Body: ").italic = True
    p_api9.add_run("{\"name\": \"Tên cập nhật\", \"description\": \"Mô tả mới\"}\n")
    p_api9.add_run("4. Xóa tâm trạng:\n").bold = True
    p_api9.add_run("• Method / Path: ").italic = True
    p_api9.add_run("DELETE /crbt-community-library/library/moods/{id}\n")
    p_api9.add_run("• Lưu ý: ").bold = True
    p_api9.add_run("Không cho phép xóa nếu tâm trạng này đang được sử dụng bởi bài nhạc nào trong hệ thống.\n")

    # Save document
    output_path = "d:/Microservice-Platform/cms_admin_api_integration_guide.docx"
    doc.save(output_path)
    print(f"Generated DOCX successfully at: {output_path}")

if __name__ == "__main__":
    main()
