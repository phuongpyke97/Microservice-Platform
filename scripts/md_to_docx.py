"""Generic Markdown -> DOCX converter for the operations docs.

Handles: ATX headings (#..####), GitHub tables, fenced code blocks (```),
bullet lists (-/*), numbered lists, blockquotes (>), inline **bold** and
`code`. Good enough for the handover docs in docs/operations/.

Usage:
    python scripts/md_to_docx.py <input.md> [<input2.md> ...]
Output: same path with .docx extension.
"""
import re
import sys

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x36, 0x5D)
SLATE = RGBColor(0x5C, 0x76, 0x8D)
CODE_BG = "F2F2F2"
INLINE_CODE = RGBColor(0xB0, 0x30, 0x60)

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def set_cell_bg(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def add_inline(paragraph, text):
    """Render **bold** and `code` spans inside a paragraph."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = INLINE_CODE
        else:
            paragraph.add_run(part)


def add_heading(doc, level, text):
    h = doc.add_heading(level=min(level, 4))
    run = h.add_run(text)
    run.font.name = "Segoe UI"
    run.font.color.rgb = NAVY if level <= 2 else SLATE
    sizes = {1: 18, 2: 15, 3: 13, 4: 11}
    run.font.size = Pt(sizes.get(level, 11))


def set_code_bg(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}"/>'))


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    set_code_bg(p)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def add_table(doc, rows):
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, cell_text in enumerate(header):
        c = table.rows[0].cells[i]
        c.paragraphs[0].clear()
        add_inline(c.paragraphs[0], cell_text)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(c, "1B365D")
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            txt = row[i] if i < len(row) else ""
            cells[i].paragraphs[0].clear()
            add_inline(cells[i].paragraphs[0], txt)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = docx.Document()
    style = doc.styles["Normal"].font
    style.name = "Segoe UI"
    style.size = Pt(10.5)
    style.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.8)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.lstrip().startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].lstrip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < n and is_separator_row(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        m = HEADING_RE.match(line)
        if m:
            add_heading(doc, len(m.group(1)), m.group(2).strip())
            i += 1
            continue

        if re.match(r"^---+$", line.strip()):
            i += 1
            continue

        if line.lstrip().startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, line.lstrip()[1:].strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = SLATE
            i += 1
            continue

        bm = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bm:
            indent = len(bm.group(1))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent // 2))
            add_inline(p, bm.group(2))
            i += 1
            continue

        nm = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, nm.group(2))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(docx_path)
    print(f"OK -> {docx_path}")


def main():
    targets = sys.argv[1:]
    if not targets:
        print("Usage: python md_to_docx.py <input.md> ...")
        sys.exit(1)
    for md in targets:
        out = re.sub(r"\.md$", ".docx", md)
        convert(md, out)


if __name__ == "__main__":
    main()
